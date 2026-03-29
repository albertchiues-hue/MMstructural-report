#!/usr/bin/env python3
"""Create mammography abnormal case report form (健康署婦女乳房X光攝影檢查服務異常個案報告表)"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Page setup: A4, narrow margins ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(0.8)
section.bottom_margin = Cm(0.6)
section.left_margin = Cm(1.0)
section.right_margin = Cm(1.0)

# ── Style defaults ──
style = doc.styles['Normal']
style.font.name = '標楷體'
style.font.size = Pt(11)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = Pt(15)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="標楷體"/>')
    rPr.append(rFonts)
else:
    rFonts.set(qn('w:eastAsia'), '標楷體')


# ── Utility functions ──
def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    for edge, val in kwargs.items():
        el = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{val}" w:space="0" w:color="000000"/>')
        existing = tcBorders.find(qn(f'w:{edge}'))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(el)

def set_all_borders(cell, sz='4'):
    set_cell_border(cell, top=sz, bottom=sz, left=sz, right=sz)

def add_run(paragraph, text, bold=False, size=None):
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = size
    run.font.name = '標楷體'
    rPr = run._r.get_or_add_rPr()
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="標楷體"/>')
        rPr.append(rF)
    else:
        rF.set(qn('w:eastAsia'), '標楷體')
    return run

def set_cell_width(cell, width_cm):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    w = str(int(width_cm * 567))
    if tcW is None:
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{w}" w:type="dxa"/>')
        tcPr.append(tcW)
    else:
        tcW.set(qn('w:w'), w)
        tcW.set(qn('w:type'), 'dxa')

def merge_cells(table, row, c1, c2):
    return table.cell(row, c1).merge(table.cell(row, c2))

def set_row_height(row, h_cm):
    trPr = row._tr.get_or_add_trPr()
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(h_cm * 567)}" w:hRule="atLeast"/>'))

def set_valign(cell, val='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:vAlign'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val}"/>'))

def cell_text(cell, text='', align=None, bold=False, size=None):
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if text:
        add_run(p, text, bold=bold, size=size)
    return p

CB = '☐'
CENTER = WD_ALIGN_PARAGRAPH.CENTER

# ══════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = CENTER
title.paragraph_format.space_after = Pt(4)
add_run(title, '健康署婦女乳房X光攝影檢查服務異常個案報告表', bold=True, size=Pt(16))

p_mrn = doc.add_paragraph()
p_mrn.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_mrn.paragraph_format.space_after = Pt(2)
add_run(p_mrn, '病歷號：', size=Pt(12))

# ══════════════════════════════════════════════
# SECTION 1: 檢查資訊 - No vertical merges to avoid issues
# Use 13 columns: col0=label, col1=field, col2=field, col3=IDlabel, col4-13=digits
# ══════════════════════════════════════════════
NC = 13
info = doc.add_table(rows=5, cols=NC)
info.alignment = WD_TABLE_ALIGNMENT.CENTER

for r in info.rows:
    for c in r.cells:
        set_all_borders(c)
        set_valign(c)

# Row 0: Header
h = merge_cells(info, 0, 0, NC - 1)
cell_text(h, '檢查資訊', align=CENTER, bold=True, size=Pt(13))
set_cell_shading(h, 'D9D9D9')
set_row_height(info.rows[0], 0.65)

# Row 1: 姓名 | (name) | 身份證統一編號 | 10 digit cells
cell_text(info.cell(1, 0), '姓名', align=CENTER, bold=True, size=Pt(11))
name1 = merge_cells(info, 1, 1, 2)
cell_text(name1, '')
cell_text(info.cell(1, 3), '身份證統一編號', align=CENTER, bold=True, size=Pt(10))
for i in range(4, NC):
    cell_text(info.cell(1, i), '', align=CENTER)
set_row_height(info.rows[1], 0.65)

# Row 2: (empty/姓名 cont) | (name cont) | 統一證號(外籍) | 10 digit cells
cell_text(info.cell(2, 0), '', align=CENTER)
name2 = merge_cells(info, 2, 1, 2)
cell_text(name2, '')
cell_text(info.cell(2, 3), '統一證號(外籍)', align=CENTER, bold=True, size=Pt(9))
for i in range(4, NC):
    cell_text(info.cell(2, i), '', align=CENTER)
set_row_height(info.rows[2], 0.65)

# Now merge vertically: 姓名 (row1,col0 + row2,col0) and name field
info.cell(1, 0).merge(info.cell(2, 0))
cell_text(info.cell(1, 0), '姓名', align=CENTER, bold=True, size=Pt(11))
info.cell(1, 1).merge(info.cell(2, 1))

# Row 3: 出生日期 | 年月日 | 攝影日期 | 年月日
cell_text(info.cell(3, 0), '出生日期', align=CENTER, bold=True, size=Pt(11))
dob = merge_cells(info, 3, 1, 5)
p = cell_text(dob, '', align=CENTER)
add_run(p, '      年      月      日', size=Pt(11))
photo_lbl = merge_cells(info, 3, 6, 7)
cell_text(photo_lbl, '攝影日期', align=CENTER, bold=True, size=Pt(11))
photo_val = merge_cells(info, 3, 8, NC - 1)
p = cell_text(photo_val, '', align=CENTER)
add_run(p, '      年      月      日', size=Pt(11))
set_row_height(info.rows[3], 0.65)

# Row 4: 醫院名稱 | field | 放射科醫師 | field
cell_text(info.cell(4, 0), '醫院名稱', align=CENTER, bold=True, size=Pt(11))
hosp = merge_cells(info, 4, 1, 5)
cell_text(hosp, '')
doc_lbl = merge_cells(info, 4, 6, 7)
cell_text(doc_lbl, '放射科醫師', align=CENTER, bold=True, size=Pt(11))
doc_val = merge_cells(info, 4, 8, NC - 1)
cell_text(doc_val, '')
set_row_height(info.rows[4], 0.65)

# ══════════════════════════════════════════════
# SECTION 2: 乳房X光攝影陽性結果 (1-col table, 2 rows)
# ══════════════════════════════════════════════
spacer = doc.add_paragraph()
spacer.paragraph_format.space_before = Pt(3)
spacer.paragraph_format.space_after = Pt(1)

bt = doc.add_table(rows=2, cols=1)
bt.alignment = WD_TABLE_ALIGNMENT.CENTER
for r in bt.rows:
    for c in r.cells:
        set_all_borders(c, '6')

# Header row
hdr = bt.cell(0, 0)
cell_text(hdr, '乳房X光攝影陽性結果', align=CENTER, bold=True, size=Pt(13))
set_cell_shading(hdr, 'D9D9D9')

# Content row
cc = bt.cell(1, 0)
p = cc.paragraphs[0]
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(1)
add_run(p, f'{CB}  Category 0：Need Additional Imaging Evaluation.', size=Pt(11))

for txt in [
    f'{CB}  Category 3：Probably Benign Finding – Short Interval Follow-up Is Suggested.',
    f'{CB}  Category 4：Suspicious Abnormality– Biopsy Should Be Considered.',
    f'        {CB} a. Low suspicion；{CB} b. Moderate suspicion；{CB} c. High suspicion；',
    f'{CB}  Category 5：Highly Suggestive of Malignancy – Appropriate Action Should Be Taken.',
]:
    np = cc.add_paragraph()
    np.paragraph_format.space_before = Pt(1)
    np.paragraph_format.space_after = Pt(1)
    add_run(np, txt, size=Pt(11))

# ══════════════════════════════════════════════
# SECTION 3: 病灶勾選 header
# ══════════════════════════════════════════════
lt = doc.add_table(rows=1, cols=1)
lt.alignment = WD_TABLE_ALIGNMENT.CENTER
lc = lt.cell(0, 0)
set_all_borders(lc, '6')
set_cell_shading(lc, 'D9D9D9')
cell_text(lc, '病灶勾選(如單側多處病灶或兩側皆有病灶，請以不同表單分開呈現)', align=CENTER, bold=True, size=Pt(11))

# ══════════════════════════════════════════════
# Helper: Location table
# ══════════════════════════════════════════════
def add_location_table():
    t = doc.add_table(rows=2, cols=7)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r in t.rows:
        for c in r.cells:
            set_all_borders(c)
            set_valign(c)
    set_row_height(t.rows[0], 0.55)
    set_row_height(t.rows[1], 0.55)

    # Row 0
    cell_text(t.cell(0, 0), 'Location', align=CENTER, bold=True, size=Pt(10))
    set_cell_shading(t.cell(0, 0), 'E8E8E8')
    for i, q in enumerate([f'{CB} UOQ', f'{CB} UIQ', f'{CB} LOQ', f'{CB} LIQ', f'{CB} Subareolar', f'{CB} Axillary tail']):
        cell_text(t.cell(0, i+1), q, align=CENTER, size=Pt(10))

    # Row 1
    cell_text(t.cell(1, 0), 'One view only', align=CENTER, size=Pt(9))
    set_cell_shading(t.cell(1, 0), 'E8E8E8')
    h1 = merge_cells(t, 1, 1, 3)
    cell_text(h1, f'{CB} Upper Hemisphere   {CB} Lower Hemisphere', align=CENTER, size=Pt(10))
    h2 = merge_cells(t, 1, 4, 6)
    cell_text(h2, f'{CB} Outer Hemisphere   {CB} Inner Hemisphere', align=CENTER, size=Pt(10))


# ══════════════════════════════════════════════
# Helper: Properties rows
# ══════════════════════════════════════════════
def add_props(props_list):
    for label, options in props_list:
        t = doc.add_table(rows=1, cols=len(options) + 1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        lbl = t.cell(0, 0)
        set_all_borders(lbl)
        set_cell_shading(lbl, 'E8E8E8')
        set_valign(lbl)
        cell_text(lbl, label, align=CENTER, bold=True, size=Pt(10))
        set_cell_width(lbl, 2.5)
        for i, opt in enumerate(options):
            c = t.cell(0, i + 1)
            set_all_borders(c)
            set_valign(c)
            cell_text(c, f'{CB} {opt}', align=CENTER, size=Pt(10))


# ══════════════════════════════════════════════
# 1. Mass
# ══════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(1)
add_run(p, f'{CB} 1. Mass：{CB} Rt.  {CB} Lt.  {CB} Multiple, Unilateral  {CB} Multiple, Bilateral', size=Pt(11))

add_location_table()
add_props([
    ('Size', ['<1.0 cm', '1-2 cm', '2-3 cm', '3-4 cm', '>4 cm']),
    ('Shape', ['Round', 'Oval', 'Lobular', 'Irregular']),
    ('Margin', ['Circumscribed', 'Microlobulated', 'Obscured', 'Indistinct', 'Spiculated']),
    ('Density', ['High density', 'Equal density', 'Low-density', 'Fat-containing']),
])

# ══════════════════════════════════════════════
# 2. Calcifications
# ══════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(1)
add_run(p, f'{CB} 2. Calcifications：{CB} Rt.  {CB} Lt.  {CB} Multiple, Unilateral  {CB} Multiple, Bilateral', size=Pt(11))

add_location_table()
add_props([
    ('Distribution', ['Grouped', 'Linear', 'Segmental', 'Regional', 'Diffuse']),
    ('Morphology', ['Amorphous', 'Coarse Heterogeneous', 'Fine Pleomorphic', 'Fine Linear Branching']),
])

# ══════════════════════════════════════════════
# 3. Asymmetry
# ══════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(1)
add_run(p, f'{CB} 3. Asymmetry：{CB} Rt.  {CB} Lt.  {CB} Asymmetry  {CB} Focal asymmetry  {CB} Developing asymmetry', size=Pt(11))

add_location_table()

# ══════════════════════════════════════════════
# 4. Architectural Distortion
# ══════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(1)
add_run(p, f'{CB} 4. Architectural Distortion：{CB} Rt.  {CB} Lt.', size=Pt(11))

add_location_table()

# ══════════════════════════════════════════════
# 5-9. Additional findings
# ══════════════════════════════════════════════
for line in [
    f'{CB}  5. Thickening or retraction of the skin and/or nipple: {CB} Rt. / {CB} Lt.',
    f'{CB}  6. Dense or enlarged axillary LNs: {CB} Rt. / {CB} Lt.',
    f'{CB}  7. Dilated lactiferous ducts: {CB} Rt. / {CB} Lt.',
    f'{CB}  8. Diffuse thickening of the skin and increased density: {CB} Rt. / {CB} Lt.',
    f'{CB}  9. Others:',
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, line, size=Pt(11))

# ══════════════════════════════════════════════
# Footer
# ══════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p, '(113年1月修訂)', size=Pt(9))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(2)
add_run(p, '第一聯：存病歷　　第二聯：存放射線科', size=Pt(9))

# ══════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════
output_path = '/Users/chiutsecheng/MMstructural report/mammography_report_form.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
