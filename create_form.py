#!/usr/bin/env python3
"""
Create mammography report form (健康署婦女乳房X光攝影檢查服務異常個案報告表)
as an A4 Word document replicating the standard Taiwan HPA form.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

FONT_NAME = "標楷體"
FONT_NAME_ASCII = "Times New Roman"
CB = "☐"  # checkbox character


def set_font(run, size=9, bold=False, font_name=FONT_NAME):
    """Set font properties on a run, including East Asian font."""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME_ASCII
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), FONT_NAME_ASCII)
    rFonts.set(qn('w:hAnsi'), FONT_NAME_ASCII)


def add_text(cell, text, size=9, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Add text to a cell, clearing existing content."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(size + 3)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return run


def shade_cell(cell, color="C0C0C0"):
    """Apply shading to a cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right with dict of sz, val, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        elem = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        existing = tcBorders.find(qn(f'w:{edge}'))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(elem)


def set_cell_width(cell, width_cm):
    """Set exact cell width."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = parse_xml(f'<w:tcW {nsdecls("w")}/>')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_cm * 567)))  # cm to twips
    tcW.set(qn('w:type'), 'dxa')


def set_row_height(row, height_cm):
    """Set exact row height."""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = parse_xml(f'<w:trPr {nsdecls("w")}></w:trPr>')
        tr.insert(0, trPr)
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_cm * 567)}" w:hRule="atLeast"/>'
    )
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


def merge_cells(table, row_start, col_start, row_end, col_end):
    """Merge a rectangle of cells."""
    start = table.cell(row_start, col_start)
    end = table.cell(row_end, col_end)
    start.merge(end)
    return table.cell(row_start, col_start)


def set_vertical_alignment(cell, align=WD_ALIGN_VERTICAL.CENTER):
    """Set vertical alignment of cell."""
    cell.vertical_alignment = align


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------
doc = Document()

# A4 page setup
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.0)
section.bottom_margin = Cm(0.8)

# Set default font
style = doc.styles['Normal']
style.font.name = FONT_NAME_ASCII
style.font.size = Pt(9)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

# ---------------------------------------------------------------------------
# Title
# ---------------------------------------------------------------------------
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(2)
title_p.paragraph_format.space_before = Pt(0)
run = title_p.add_run("健康署婦女乳房X光攝影檢查服務異常個案報告表")
set_font(run, size=14, bold=True, font_name=FONT_NAME)

# 病歷號
mrn_p = doc.add_paragraph()
mrn_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
mrn_p.paragraph_format.space_after = Pt(2)
mrn_p.paragraph_format.space_before = Pt(0)
run = mrn_p.add_run("病歷號：")
set_font(run, size=10, bold=False)

# ---------------------------------------------------------------------------
# 檢查資訊 Table (Patient Info) — 5 rows × 12 cols
# ---------------------------------------------------------------------------
NUM_COLS_INFO = 12
info_table = doc.add_table(rows=5, cols=NUM_COLS_INFO)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set all borders
for row in info_table.rows:
    for cell in row.cells:
        set_cell_border(cell,
            top={"sz": "4", "val": "single", "color": "000000"},
            bottom={"sz": "4", "val": "single", "color": "000000"},
            left={"sz": "4", "val": "single", "color": "000000"},
            right={"sz": "4", "val": "single", "color": "000000"})

# Row 0: 檢查資訊 header
cell = merge_cells(info_table, 0, 0, 0, NUM_COLS_INFO - 1)
add_text(cell, "檢查資訊", size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
shade_cell(cell, "D9D9D9")
set_row_height(info_table.rows[0], 0.6)

# Row 1: 姓名 | 身份證統一編號 boxes | 統一證號(外籍) boxes
cell = merge_cells(info_table, 1, 0, 2, 1)
add_text(cell, "姓名", size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_vertical_alignment(cell, WD_ALIGN_VERTICAL.CENTER)

# 身份證統一編號 label + boxes
cell = merge_cells(info_table, 1, 2, 1, 6)
add_text(cell, "身份證統一編號", size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# ID number boxes (row 1, cols 7-11 as individual cells for boxes effect)
for c in range(7, NUM_COLS_INFO):
    cell = info_table.cell(1, c)
    add_text(cell, "", size=9)

# Row 2: 統一證號(外籍) label + boxes
cell = merge_cells(info_table, 2, 2, 2, 6)
add_text(cell, "統一證號(外籍)", size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

for c in range(7, NUM_COLS_INFO):
    cell = info_table.cell(2, c)
    add_text(cell, "", size=9)

set_row_height(info_table.rows[1], 0.6)
set_row_height(info_table.rows[2], 0.6)

# Row 3: 出生日期 | 攝影日期
cell = merge_cells(info_table, 3, 0, 3, 1)
add_text(cell, "出生日期", size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

cell = merge_cells(info_table, 3, 2, 3, 5)
add_text(cell, "　　年　　月　　日", size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

cell = merge_cells(info_table, 3, 6, 3, 7)
add_text(cell, "攝影日期", size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

cell = merge_cells(info_table, 3, 8, 3, NUM_COLS_INFO - 1)
add_text(cell, "　　年　　月　　日", size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_row_height(info_table.rows[3], 0.6)

# Row 4: 醫院名稱 | 放射科醫師
cell = merge_cells(info_table, 4, 0, 4, 1)
add_text(cell, "醫院名稱", size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

cell = merge_cells(info_table, 4, 2, 4, 5)
add_text(cell, "", size=9)

cell = merge_cells(info_table, 4, 6, 4, 7)
add_text(cell, "放射科醫師", size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

cell = merge_cells(info_table, 4, 8, 4, NUM_COLS_INFO - 1)
add_text(cell, "", size=9)
set_row_height(info_table.rows[4], 0.6)

# ---------------------------------------------------------------------------
# 乳房X光攝影陽性結果 header
# ---------------------------------------------------------------------------
header_p = doc.add_paragraph()
header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
header_p.paragraph_format.space_before = Pt(4)
header_p.paragraph_format.space_after = Pt(2)
run = header_p.add_run("乳房X光攝影陽性結果")
set_font(run, size=10, bold=True)

# ---------------------------------------------------------------------------
# Main findings table
# We'll build this as a large table with careful merging
# ---------------------------------------------------------------------------

# Categories section (separate paragraphs for each category)
cat_lines = [
    f"{CB} Category 0：Need Additional Imaging Evaluation.",
    f"{CB} Category 3：Probably Benign Finding – Short Interval Follow-up Is Suggested.",
    f"{CB} Category 4：Suspicious Abnormality– Biopsy Should Be Considered.",
    f"　　　{CB} a. Low suspicion；{CB} b. Moderate suspicion；{CB} c. High suspicion；",
    f"{CB} Category 5：Highly Suggestive of Malignancy – Appropriate Action Should Be Taken.",
]

for line in cat_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(13)
    run = p.add_run(line)
    set_font(run, size=9)

# ---------------------------------------------------------------------------
# 病灶勾選 instruction bar
# ---------------------------------------------------------------------------
instr_p = doc.add_paragraph()
instr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
instr_p.paragraph_format.space_before = Pt(3)
instr_p.paragraph_format.space_after = Pt(2)
# Add shading to paragraph
pPr = instr_p._element.get_or_add_pPr()
shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="C0C0C0" w:val="clear"/>')
pPr.append(shd)
run = instr_p.add_run("病灶勾選(如單側多處病灶或兩側皆有病灶，請以不同表單分開呈現)")
set_font(run, size=9, bold=True)

# ---------------------------------------------------------------------------
# Findings Table — complex multi-section table
# We use a single large table for the entire findings section
# Columns: col0 (label ~2cm), col1-col5 (content, ~3cm each)
# ---------------------------------------------------------------------------

NUM_COLS = 6

# Build rows programmatically
findings_table = doc.add_table(rows=0, cols=NUM_COLS)
findings_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Helper to add a row
def add_row():
    row = findings_table.add_row()
    for cell in row.cells:
        set_cell_border(cell,
            top={"sz": "4", "val": "single", "color": "000000"},
            bottom={"sz": "4", "val": "single", "color": "000000"},
            left={"sz": "4", "val": "single", "color": "000000"},
            right={"sz": "4", "val": "single", "color": "000000"})
    return row


def add_section_header(text, shade=False):
    """Add a full-width header row."""
    row = add_row()
    cell = merge_cells(findings_table, len(findings_table.rows)-1, 0,
                       len(findings_table.rows)-1, NUM_COLS-1)
    add_text(cell, text, size=9, bold=False)
    if shade:
        shade_cell(cell, "D9D9D9")
    set_row_height(row, 0.55)
    return row


def add_location_rows(row_offset):
    """Add the standard location sub-rows (quadrant + hemisphere)."""
    # Row: quadrant line
    row = add_row()
    ri = len(findings_table.rows) - 1
    cell = merge_cells(findings_table, ri, 0, ri, 0)
    add_text(cell, "Location", size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_vertical_alignment(cell, WD_ALIGN_VERTICAL.CENTER)

    cell = merge_cells(findings_table, ri, 1, ri, NUM_COLS-1)
    add_text(cell, f"  {CB} UOQ　{CB} UIQ　{CB} LOQ　{CB} LIQ　{CB} Subareolar　{CB} Axillary tail", size=8)
    set_row_height(row, 0.5)

    # Row: hemisphere options (with "One view only")
    row = add_row()
    ri = len(findings_table.rows) - 1
    # Merge location label with row above
    merge_cells(findings_table, ri-1, 0, ri, 0)
    cell = findings_table.cell(ri-1, 0)
    # Re-add text after merge
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Location")
    set_font(run, size=8, bold=True)
    set_vertical_alignment(cell, WD_ALIGN_VERTICAL.CENTER)

    cell = merge_cells(findings_table, ri, 1, ri, 1)
    add_text(cell, "One view only", size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    cell = merge_cells(findings_table, ri, 2, ri, 3)
    add_text(cell, f"  {CB} Upper Hemisphere　{CB} Lower Hemisphere", size=8)

    cell = merge_cells(findings_table, ri, 4, ri, NUM_COLS-1)
    add_text(cell, f"  {CB} Outer Hemisphere　{CB} Inner Hemisphere", size=8)
    set_row_height(row, 0.5)


def add_property_row(label, options_text):
    """Add a single property row (e.g., Size, Shape, Margin, Density)."""
    row = add_row()
    ri = len(findings_table.rows) - 1
    cell = findings_table.cell(ri, 0)
    add_text(cell, label, size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_vertical_alignment(cell, WD_ALIGN_VERTICAL.CENTER)
    shade_cell(cell, "D9D9D9")

    cell = merge_cells(findings_table, ri, 1, ri, NUM_COLS-1)
    add_text(cell, options_text, size=8)
    set_row_height(row, 0.5)
    return row


# ===== 1. Mass =====
add_section_header(
    f"{CB} 1. Mass：{CB} Rt.  {CB} Lt.  {CB} Multiple, Unilateral  {CB} Multiple, Bilateral"
)

add_location_rows(0)

add_property_row("Size",
    f"  {CB} <1.0 cm　　{CB} 1-2 cm　　{CB} 2-3 cm　　{CB} 3-4 cm　　{CB} > 4 cm")

add_property_row("Shape",
    f"  {CB} Round　　{CB} Oval　　{CB} Lobular　　{CB} Irregular")

add_property_row("Margin",
    f"  {CB} Circumscribed　{CB} Microlobulated　{CB} Obscured　{CB} Indistinct　{CB} Spiculated")

add_property_row("Density",
    f"  {CB} High density　{CB} Equal density　{CB} Low-density　{CB} Fat-containing")

# ===== 2. Calcifications =====
add_section_header(
    f"{CB} 2. Calcifications：{CB} Rt.  {CB} Lt.  {CB} Multiple, Unilateral  {CB} Multiple, Bilateral"
)

add_location_rows(0)

add_property_row("Distribution",
    f"  {CB} Grouped　{CB} Linear　{CB} Segmental　{CB} Regional　{CB} Diffuse")

add_property_row("Morphology",
    f"  {CB} Amorphous  {CB} Coarse Heterogeneous  {CB} Fine Pleomorphic  {CB} Fine Linear Branching")

# ===== 3. Asymmetry =====
add_section_header(
    f"{CB} 3. Asymmetry：{CB} Rt.  {CB} Lt.  {CB} Asymmetry  {CB} Focal asymmetry  {CB} Developing asymmetry"
)

add_location_rows(0)

# ===== 4. Architectural Distortion =====
add_section_header(
    f"{CB} 4. Architectural Distortion：{CB} Rt.  {CB} Lt."
)

add_location_rows(0)

# ===== 5-9 Additional findings =====
add_section_header(
    f"{CB} 5. Thickening or retraction of the skin and/or nipple：{CB} Rt. / {CB} Lt."
)
add_section_header(
    f"{CB} 6. Dense or enlarged axillary LNs：{CB} Rt. / {CB} Lt."
)
add_section_header(
    f"{CB} 7. Dilated lactiferous ducts：{CB} Rt. / {CB} Lt."
)
add_section_header(
    f"{CB} 8. Diffuse thickening of the skin and increased density：{CB} Rt. / {CB} Lt."
)
add_section_header(
    f"{CB} 9. Others："
)

# ---------------------------------------------------------------------------
# Footer note
# ---------------------------------------------------------------------------
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_p.paragraph_format.space_before = Pt(2)
footer_p.paragraph_format.space_after = Pt(0)
run = footer_p.add_run("(113 年 1 月修訂)")
set_font(run, size=7)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
output_path = "/Users/chiutsecheng/MMstructural report/mammography_report_form.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
